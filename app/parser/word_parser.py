# word_parser.py
from docx import Document
from pathlib import Path


def parse_word(file_path: Path) -> str:
    """解析Word文档，返回文本内容"""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n\n".join(text_parts)
